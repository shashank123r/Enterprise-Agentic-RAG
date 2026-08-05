"""DOCX document extractor — sync calls offloaded via run_in_executor."""

from pathlib import Path

from app.core.logging import get_logger
from app.ingestion.executor import run_in_executor
from app.ingestion.extractors import (
    BaseExtractor,
    ExtractionResult,
    ImageResult,
    PageResult,
    TableResult,
    extractor_registry,
)

logger = get_logger(__name__)

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@extractor_registry.register(_DOCX_MIME)
class DocxExtractor(BaseExtractor):
    """Extract text, tables, and images from DOCX documents."""

    async def extract(self) -> ExtractionResult:
        def _do_extract() -> ExtractionResult:
            from docx import Document as DocxDocument
            from docx.opc.constants import RELATIONSHIP_TYPE as RT

            r = ExtractionResult()
            doc = DocxDocument(str(self.file_path))
            props = doc.core_properties
            r.metadata.update(
                {
                    "title": props.title or "",
                    "author": props.author or "",
                    "created": str(props.created) if props.created else "",
                    "modified": str(props.modified) if props.modified else "",
                    "last_modified_by": props.last_modified_by or "",
                    "category": props.category or "",
                    "keywords": props.keywords or "",
                    "revision": props.revision or 0,
                    "page_count": len(doc.paragraphs),
                }
            )

            page_text: list[str] = []
            current_section: str | None = None
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                if para.style and "Heading" in (para.style.name or ""):
                    if page_text:
                        r.pages.append(
                            PageResult(
                                page_number=len(r.pages) + 1,
                                text="\n".join(page_text),
                                section_title=current_section,
                            )
                        )
                        page_text = []
                    current_section = text
                page_text.append(text)
            if page_text:
                r.pages.append(
                    PageResult(
                        page_number=len(r.pages) + 1,
                        text="\n".join(page_text),
                        section_title=current_section,
                    )
                )

            for t_idx, table in enumerate(doc.tables):
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                rows = [[cell.text.strip() for cell in row.cells] for row in table.rows[1:]]
                r.tables.append(
                    TableResult(page_number=None, table_index=t_idx, headers=headers, rows=rows)
                )

            image_count = 0
            for rel in doc.part.rels.values():
                if rel.reltype == RT.IMAGE:
                    image_count += 1
                    ext = Path(rel.target_ref or "").suffix.lstrip(".").lower() or "png"
                    r.images.append(
                        ImageResult(
                            page_number=None,
                            image_index=image_count,
                            image_data=rel.target_part.blob,
                            format=ext,
                        )
                    )

            r.text = "\n\n".join(p.text for p in r.pages if p.text.strip())
            return r

        return await run_in_executor(_do_extract)

    async def extract_text(self) -> str:
        def _do_text() -> str:
            from docx import Document as DocxDocument

            doc = DocxDocument(str(self.file_path))
            return "\n".join(p.text for p in doc.paragraphs)

        return await run_in_executor(_do_text)

    async def extract_metadata(self) -> dict:
        def _do_meta() -> dict:
            from docx import Document as DocxDocument

            doc = DocxDocument(str(self.file_path))
            props = doc.core_properties
            return {
                "title": props.title or "",
                "author": props.author or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
                "last_modified_by": props.last_modified_by or "",
                "category": props.category or "",
                "keywords": props.keywords or "",
                "revision": props.revision or 0,
                "page_count": len(doc.paragraphs),
            }

        return await run_in_executor(_do_meta)
